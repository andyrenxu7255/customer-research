#!/usr/bin/env python3
"""
生成客户调研 Word 报告
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from datetime import datetime

def create_word_report(customer_name, results, output_path):
    """创建 Word 调研报告"""
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'客户调研报告：{customer_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f'**调研时间：** {datetime.now().strftime("%Y 年 %m 月 %d 日")}')
    doc.add_paragraph(f'**调研人：** Dick Dunkel')
    doc.add_paragraph(f'**客户类型：** 民营医疗机构')
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('📌 执行摘要', level=1)
    doc.add_paragraph(f'本次调研围绕 {customer_name} 的 AI 相关领导发言、招标记录、数据相关内容及基本情况展开。')
    doc.add_paragraph()
    
    # 各维度详情
    sections = [
        ("1️⃣ AI 相关领导发言", "AI 相关领导发言"),
        ("2️⃣ AI 相关招标记录", "AI 相关招标记录"),
        ("3️⃣ 数据相关内容", "数据相关内容"),
        ("4️⃣ 基本情况与最新动态", "基本情况与最新动态")
    ]
    
    for emoji_name, key_name in sections:
        doc.add_heading(emoji_name, level=2)
        items = results.get(key_name, [])
        if items:
            for i, item in enumerate(items, 1):
                p = doc.add_paragraph(style='List Number')
                # 清理文本，只保留前 300 字符
                clean_text = item[:300].replace('\n', ' ').strip()
                p.add_run(f'{clean_text}...')
        else:
            doc.add_paragraph('暂无相关公开信息', style='Intense Quote')
        doc.add_paragraph()
    
    # 销售机会分析
    doc.add_heading('🎯 销售机会分析', level=1)
    doc.add_paragraph('可能的痛点：', style='Heading 2')
    pain_points = [
        '多院区数据孤岛问题',
        '业务多元化需要综合医疗信息系统',
        '数字化转型压力',
        '三级医院升级需要配套信息化系统'
    ]
    for point in pain_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph('决策链推测：', style='Heading 2')
    doc.add_paragraph('• EB（经济决策人）：胡澜（CEO）或字节跳动管理层')
    doc.add_paragraph('• Technical Buyer：信息科主任/CTO（待拜访确认）')
    doc.add_paragraph('• Champion：待识别')
    
    doc.add_paragraph('切入建议：', style='Heading 2')
    suggestions = [
        '首访目标：信息科/IT 部门负责人',
        '切入点：多院区数据治理、智能问数、3D 可视化',
        'MEDDIC 行动：确认预算、决策流程、培养 Champion'
    ]
    for sug in suggestions:
        doc.add_paragraph(sug, style='List Bullet')
    
    # 信息来源
    doc.add_heading('📎 信息来源', level=1)
    doc.add_paragraph('本报告基于公开网络搜索整理，主要来源包括：')
    doc.add_paragraph('• 百度百科', style='List Bullet')
    doc.add_paragraph('• Bing 搜索引擎', style='List Bullet')
    doc.add_paragraph('• 公开新闻报道', style='List Bullet')
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ Word 报告已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法：python create_word.py <客户名称> <输出路径>")
        sys.exit(1)
    
    customer_name = sys.argv[1]
    output_path = sys.argv[2]
    
    # 示例数据（实际应从调研结果传入）
    results = {
        "AI 相关领导发言": ["公开渠道未检索到 AI 相关领导发言"],
        "AI 相关招标记录": ["民营医院可能不公开招标，需直接询问"],
        "数据相关内容": ["推测有多院区数据治理、EMR、HIS 等需求"],
        "基本情况与最新动态": ["字节跳动 100% 控股，2024 年完成向综合医疗平台转型"]
    }
    
    create_word_report(customer_name, results, output_path)
